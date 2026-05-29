
from docx import Document


def create_docx(text, output_path):

    document = Document()

    document.add_heading(

        "MeetNote AI Notes",

        level=1

    )

    document.add_paragraph(text)

    document.save(output_path)